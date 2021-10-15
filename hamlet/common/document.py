from chardet.universaldetector import UniversalDetector
from django.core.exceptions import ValidationError
import docx
import magic


def factory(fp):
    """Factory for creating a document object.

    A Document must implement a `words` property that consists of a list
    of words in that document.

    :param fp: `django.core.files.uploadedfile.UploadedFile`
    :return: document object
    """
    mimetype = magic.from_buffer(fp.read(8192), mime=True)
    if mimetype in ["text/plain", "text/x-c"]:
        return TextDocument(fp)
    elif mimetype in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                      "application/zip",
                      "application/msword"]:
        # magic has been observed identifying docx as both mimetypes
        try:
            return DocxDocument(fp)
        except:
            raise ValidationError(f"Invalid document type: {mimetype}")
    else:
        raise ValidationError(f"Invalid document type: {mimetype}")


class TextDocument:
    """Document object for representing a text document."""
    def __init__(self, doc):
        self.doc = doc
        self._words = []
        self._encoding = ""

    @property
    def words(self):
        if not self._words:
            for line in self.doc:
                self._words.extend(line.decode(self.encoding).strip().split())
        return self._words

    @property
    def encoding(self):
        if not self._encoding:
            detector = UniversalDetector()
            for line in self.doc:
                detector.feed(line)
                if detector.done:
                    break
            detector.close()
            self._encoding = detector.result['encoding']
        return self._encoding


class DocxDocument:
    """Document object for representing a DOCX document."""
    def __init__(self, doc):
        self._words = []
        self.doc = docx.Document(doc)

    @property
    def words(self):
        if not self._words:
            for p in self.doc.paragraphs:
                self._words.extend(p.text.strip().split())
        return self._words
